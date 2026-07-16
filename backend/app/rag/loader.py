"""Document loaders that normalize supported business files into content units."""
import csv
import io
import re
from html import unescape
from pathlib import Path
from typing import Any, Dict, List, Optional

import fitz  # PyMuPDF

from app.utils.files import extension_for


TABLE_ROW_WINDOW = 25


class PDFLoader:
    """Backward-compatible PDF loader used by tests and PDF ingestion."""

    @staticmethod
    def load_pdf(file_path: str) -> Dict[str, Any]:
        try:
            doc = fitz.open(file_path)
            filename = Path(file_path).name
            pages = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                pages.append(
                    {
                        "page_num": page_num + 1,
                        "text": page.get_text(),
                        "metadata": {
                            "width": page.rect.width,
                            "height": page.rect.height,
                        },
                    }
                )
            doc.close()
            return {"filename": filename, "pages": pages, "total_pages": len(pages)}
        except Exception as exc:
            raise Exception(f"Failed to load PDF {file_path}: {str(exc)}") from exc

    @staticmethod
    def validate_pdf(file_path: str) -> bool:
        try:
            doc = fitz.open(file_path)
            is_valid = len(doc) > 0
            doc.close()
            return is_valid
        except Exception:
            return False


class DocumentLoaderRegistry:
    """Load supported files into normalized units with citation locators."""

    def load_document(self, file_path: str, original_filename: Optional[str] = None) -> Dict[str, Any]:
        extension = extension_for(original_filename or file_path)
        filename = original_filename or Path(file_path).name
        if extension == ".pdf":
            return _load_pdf_document(file_path, filename)
        if extension == ".docx":
            return _load_docx_document(file_path, filename)
        if extension in {".txt", ".md"}:
            return _load_text_document(file_path, filename, extension)
        if extension in {".html", ".htm"}:
            return _load_html_document(file_path, filename)
        if extension == ".csv":
            return _load_csv_document(file_path, filename)
        if extension == ".xlsx":
            return _load_xlsx_document(file_path, filename)
        raise ValueError(f"Unsupported file type: {extension or 'none'}")


def _load_pdf_document(file_path: str, filename: str) -> Dict[str, Any]:
    pdf_data = PDFLoader.load_pdf(file_path)
    units = [
        {
            "unit_num": page["page_num"],
            "page_num": page["page_num"],
            "locator": f"page:{page['page_num']}",
            "locator_label": f"Page {page['page_num']}",
            "section_title": None,
            "sheet_name": None,
            "text": page["text"],
            "metadata": page.get("metadata", {}),
        }
        for page in pdf_data["pages"]
    ]
    return _document_payload(filename, "pdf", units)


def _load_docx_document(file_path: str, filename: str) -> Dict[str, Any]:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise ValueError("DOCX support requires python-docx to be installed") from exc

    doc = DocxDocument(file_path)
    sections = []
    current_title = None
    current_lines: List[str] = []

    def flush() -> None:
        nonlocal current_lines, current_title
        text = "\n".join(line for line in current_lines if line.strip()).strip()
        if text:
            sections.append((current_title, text))
        current_lines = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        if style_name.startswith("heading"):
            flush()
            current_title = text
        else:
            current_lines.append(text)
    flush()

    for table_index, table in enumerate(doc.tables, start=1):
        lines = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
        if lines:
            sections.append((f"Table {table_index}", "\n".join(lines)))

    if not sections:
        joined = "\n".join(paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip())
        sections = [(None, joined)] if joined else []
    return _document_payload(filename, "docx", _section_units(sections))


def _load_text_document(file_path: str, filename: str, extension: str) -> Dict[str, Any]:
    text = _read_text(file_path)
    sections = _split_markdown_sections(text) if extension == ".md" else _split_plain_text_sections(text)
    return _document_payload(filename, extension.lstrip("."), _section_units(sections))


def _load_html_document(file_path: str, filename: str) -> Dict[str, Any]:
    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise ValueError("HTML support requires beautifulsoup4 to be installed") from exc

    soup = BeautifulSoup(_read_text(file_path), "html.parser")
    for selector in ["script", "style", "nav", "header", "footer", "aside", "form", "svg"]:
        for node in soup.select(selector):
            node.decompose()
    sections = []
    current_title = None
    current_lines: List[str] = []

    def flush() -> None:
        nonlocal current_title, current_lines
        text = "\n".join(current_lines).strip()
        if text:
            sections.append((current_title, text))
        current_lines = []

    body = soup.body or soup
    for node in body.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "th"], recursive=True):
        text = " ".join(node.get_text(" ", strip=True).split())
        if not text:
            continue
        if re.fullmatch(r"h[1-6]", node.name or ""):
            flush()
            current_title = unescape(text)
        else:
            current_lines.append(unescape(text))
    flush()
    if not sections:
        text = " ".join(body.get_text(" ", strip=True).split())
        sections = [(None, unescape(text))] if text else []
    return _document_payload(filename, "html", _section_units(sections))


def _load_csv_document(file_path: str, filename: str) -> Dict[str, Any]:
    with open(file_path, "r", encoding="utf-8-sig", newline="") as file:
        rows = list(csv.reader(file))
    units = _table_units(rows, filename, sheet_name=None, start_row_offset=1)
    return _document_payload(filename, "csv", units)


def _load_xlsx_document(file_path: str, filename: str) -> Dict[str, Any]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise ValueError("XLSX support requires openpyxl to be installed") from exc

    workbook = load_workbook(file_path, read_only=True, data_only=True)
    units = []
    try:
        for sheet in workbook.worksheets:
            rows = [[_cell_to_text(value) for value in row] for row in sheet.iter_rows(values_only=True)]
            units.extend(_table_units(rows, filename, sheet_name=sheet.title, start_row_offset=1))
    finally:
        workbook.close()
    return _document_payload(filename, "xlsx", units)


def _document_payload(filename: str, document_type: str, units: List[Dict[str, Any]]) -> Dict[str, Any]:
    return {
        "filename": filename,
        "document_type": document_type,
        "units": units,
        "pages": [
            {"page_num": unit.get("page_num") or unit["unit_num"], "text": unit["text"]}
            for unit in units
        ],
        "total_units": len(units),
        "total_pages": len(units) if document_type == "pdf" else 0,
    }


def _section_units(sections: List[tuple[Optional[str], str]]) -> List[Dict[str, Any]]:
    units = []
    for index, (title, text) in enumerate(sections, start=1):
        label = f"Section {index}" + (f": {title}" if title else "")
        units.append(
            {
                "unit_num": index,
                "page_num": None,
                "locator": f"section:{index}",
                "locator_label": label,
                "section_title": title,
                "sheet_name": None,
                "text": text,
                "metadata": {},
            }
        )
    return units


def _table_units(
    rows: List[List[str]],
    filename: str,
    sheet_name: Optional[str],
    start_row_offset: int,
) -> List[Dict[str, Any]]:
    clean_rows = [[_cell_to_text(value) for value in row] for row in rows if any(_cell_to_text(value) for value in row)]
    if not clean_rows:
        return []
    header = clean_rows[0]
    data_rows = clean_rows[1:] or clean_rows[:1]
    row_start_base = start_row_offset + (1 if len(clean_rows) > 1 else 0)
    units = []
    for index, start in enumerate(range(0, len(data_rows), TABLE_ROW_WINDOW), start=1):
        window = data_rows[start:start + TABLE_ROW_WINDOW]
        start_row = row_start_base + start
        end_row = start_row + len(window) - 1
        prefix = f"{sheet_name}, " if sheet_name else ""
        label = f"{prefix}Rows {start_row}-{end_row}"
        lines = [f"Source: {filename}", f"Locator: {label}"]
        if header:
            lines.append("Columns: " + " | ".join(header))
        for row_number, row in zip(range(start_row, end_row + 1), window):
            lines.append(f"Row {row_number}: " + " | ".join(row))
        units.append(
            {
                "unit_num": index,
                "page_num": None,
                "locator": f"{sheet_name or 'sheet'}:rows:{start_row}-{end_row}",
                "locator_label": label,
                "section_title": None,
                "sheet_name": sheet_name,
                "text": "\n".join(lines),
                "metadata": {"row_start": start_row, "row_end": end_row},
            }
        )
    return units


def _split_markdown_sections(text: str) -> List[tuple[Optional[str], str]]:
    sections = []
    title = None
    lines: List[str] = []
    for line in text.splitlines():
        heading = re.match(r"^\s{0,3}#{1,6}\s+(.+?)\s*$", line)
        if heading:
            if lines:
                sections.append((title, "\n".join(lines).strip()))
                lines = []
            title = heading.group(1).strip()
        else:
            lines.append(line)
    if lines:
        sections.append((title, "\n".join(lines).strip()))
    return [(section_title, section_text) for section_title, section_text in sections if section_text]


def _split_plain_text_sections(text: str) -> List[tuple[Optional[str], str]]:
    blocks = [block.strip() for block in re.split(r"\n\s*\n", text) if block.strip()]
    if not blocks:
        return []
    return [(None, "\n\n".join(blocks))]


def _read_text(file_path: str) -> str:
    content = Path(file_path).read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Document text could not be decoded")


def _cell_to_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).replace("\n", " ").strip()
