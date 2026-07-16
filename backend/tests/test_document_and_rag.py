import importlib.util
import tempfile
import unittest
import zipfile
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

import fitz
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.db.models import Base, Document
from app.rag.chunker import TextChunker
from app.rag.loader import DocumentLoaderRegistry, PDFLoader
from app.rag.pipeline import RAGPipeline, _build_context
from app.services.document_service import DocumentService
from app.utils.files import expand_upload


class FakeVectorStore:
    def __init__(self):
        self.reset_calls = 0
        self.upserts = []
        self.deleted = []

    def reset_collection(self):
        self.reset_calls += 1

    def create_collection(self):
        pass

    def upsert_vectors(self, embeddings, metadatas):
        self.upserts.append((embeddings, metadatas))
        return list(range(len(embeddings)))

    def delete_document_vectors(self, document_id):
        self.deleted.append(document_id)


class DocumentAndRagTests(unittest.TestCase):
    def setUp(self):
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        self.session = sessionmaker(bind=engine)()

    def tearDown(self):
        self.session.close()

    def test_chunker_preserves_page_metadata_and_overlap(self):
        chunks = TextChunker.chunk_pages(
            [{"page_num": 3, "text": "One two three four five six seven eight nine ten"}],
            chunk_size=20,
            chunk_overlap=5,
        )
        self.assertGreater(len(chunks), 1)
        self.assertTrue(all(chunk["page_num"] == 3 for chunk in chunks))
        self.assertEqual([chunk["chunk_id"] for chunk in chunks], list(range(len(chunks))))

    def test_pdf_loader_reads_pdf_and_rejects_invalid_file(self):
        with tempfile.TemporaryDirectory() as directory:
            pdf_path = Path(directory) / "sample.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text((72, 72), "DocuQuery PDF loader test")
            document.save(pdf_path)
            document.close()

            self.assertTrue(PDFLoader.validate_pdf(str(pdf_path)))
            loaded = PDFLoader.load_pdf(str(pdf_path))
            self.assertEqual(loaded["total_pages"], 1)
            self.assertIn("DocuQuery PDF loader test", loaded["pages"][0]["text"])

        with tempfile.NamedTemporaryFile(suffix=".pdf") as invalid:
            invalid.write(b"not a pdf")
            invalid.flush()
            self.assertFalse(PDFLoader.validate_pdf(invalid.name))

    @unittest.skipUnless(
        all(importlib.util.find_spec(name) for name in ("docx", "bs4", "openpyxl")),
        "optional document parser dependencies are not installed",
    )
    def test_document_loader_registry_extracts_business_formats(self):
        from docx import Document as DocxDocument
        from openpyxl import Workbook

        with tempfile.TemporaryDirectory() as directory:
            base = Path(directory)
            docx_path = base / "policy.docx"
            docx = DocxDocument()
            docx.add_heading("Leave Policy", level=1)
            docx.add_paragraph("Employees receive 20 days of annual leave.")
            docx.save(docx_path)

            md_path = base / "handbook.md"
            md_path.write_text("# Support\nUse calm, clear customer updates.", encoding="utf-8")

            html_path = base / "support.html"
            html_path.write_text(
                "<html><body><nav>Menu</nav><h1>Ticket Updates</h1><p>Follow up daily.</p></body></html>",
                encoding="utf-8",
            )

            csv_path = base / "regions.csv"
            csv_path.write_text("region,sla\nUS,4h\nEU,8h\n", encoding="utf-8")

            xlsx_path = base / "support.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Support"
            sheet.append(["region", "sla"])
            sheet.append(["US", "4h"])
            workbook.save(xlsx_path)

            registry = DocumentLoaderRegistry()
            self.assertEqual(registry.load_document(str(docx_path), "policy.docx")["document_type"], "docx")
            self.assertIn("Leave Policy", registry.load_document(str(docx_path), "policy.docx")["units"][0]["locator_label"])
            self.assertEqual(registry.load_document(str(md_path), "handbook.md")["units"][0]["section_title"], "Support")
            self.assertNotIn("Menu", registry.load_document(str(html_path), "support.html")["units"][0]["text"])
            self.assertEqual(registry.load_document(str(csv_path), "regions.csv")["units"][0]["locator_label"], "Rows 2-3")
            self.assertEqual(registry.load_document(str(xlsx_path), "support.xlsx")["units"][0]["locator_label"], "Support, Rows 2-2")

    def test_zip_upload_partially_accepts_safe_files_and_rejects_unsafe_entries(self):
        archive = BytesIO()
        with zipfile.ZipFile(archive, "w") as zip_file:
            zip_file.writestr("docs/policy.txt", "Safe public policy")
            zip_file.writestr("../escape.txt", "bad")
            zip_file.writestr("nested/archive.zip", b"PK\x03\x04bad")
            zip_file.writestr(".DS_Store", "junk")
            zip_file.writestr("script.js", "alert(1)")

        accepted, rejected = expand_upload("mixed.zip", archive.getvalue())
        self.assertEqual([item.filename for item in accepted], ["policy.txt"])
        rejected_reasons = {item.filename: item.reason for item in rejected}
        self.assertIn("../escape.txt", rejected_reasons)
        self.assertIn("nested/archive.zip", rejected_reasons)
        self.assertIn(".DS_Store", rejected_reasons)
        self.assertIn("script.js", rejected_reasons)

    def test_upload_validation_rejects_spoofed_or_binary_content(self):
        accepted, rejected = expand_upload("fake.pdf", b"not really a pdf")
        self.assertEqual(accepted, [])
        self.assertIn("valid PDF", rejected[0].reason)

        accepted, rejected = expand_upload("binary.txt", b"\x01\x02\x03\x04" * 20)
        self.assertEqual(accepted, [])
        self.assertIn("binary", rejected[0].reason)

    def test_rag_pipeline_refuses_without_chunks_and_formats_debug_sources(self):
        pipeline = RAGPipeline.__new__(RAGPipeline)
        pipeline.retriever = type(
            "EmptyRetriever", (), {"retrieve": lambda self, *args, **kwargs: []}
        )()
        pipeline.generator = None
        refused = pipeline.run("Unknown?", show_debug=True)
        self.assertIn("could not find", refused["answer"].lower())
        self.assertEqual(refused["retrieved_chunks"], [])

        pipeline.retriever = type(
            "Retriever",
            (),
            {
                "retrieve": lambda self, *args, **kwargs: [
                    {
                        "filename": "policy.pdf",
                        "page": 1,
                        "locator_label": "Page 1",
                        "chunk_id": 7,
                        "text": "Refunds are available in 30 days.",
                        "preview": "Refunds are available",
                        "rank": 1,
                        "score": 0.9,
                    }
                ]
            },
        )()
        pipeline.generator = type(
            "Generator",
            (),
            {"generate": lambda self, question, context, prompt_variant=None: "30 days"},
        )()
        answered = pipeline.run("Refund window?", retrieval_method="hybrid", show_debug=True)
        self.assertEqual(answered["answer"], "30 days")
        self.assertEqual(answered["sources"][0]["filename"], "policy.pdf")
        self.assertEqual(answered["sources"][0]["locator_label"], "Page 1")
        self.assertIn("30 days", answered["retrieved_chunks"][0]["full_text"])

    def test_rag_pipeline_can_auto_route_and_verify_answer(self):
        calls = []
        pipeline = RAGPipeline.__new__(RAGPipeline)
        pipeline.retriever = type(
            "Retriever",
            (),
            {
                "retrieve": lambda self, question, top_k, method, reranker: (
                    calls.append((top_k, method, reranker))
                    or [
                        {
                            "filename": "policy.pdf",
                            "page": 2,
                            "locator_label": "Page 2",
                            "chunk_id": 8,
                            "text": "Employees must report violations through EthicsPoint.",
                            "preview": "Report violations through EthicsPoint",
                            "rank": 1,
                            "score": 0.9,
                        }
                    ]
                )
            },
        )()
        pipeline.generator = type(
            "Generator",
            (),
            {
                "generate": lambda self, question, context, prompt_variant=None: "Use Lighthouse.",
                "verify_answer": lambda self, question, context, draft: "Use EthicsPoint.",
            },
        )()

        result = pipeline.run(
            "How are policy violations reported?",
            retrieval_profile="auto",
            answer_verification=True,
            question_type="policy_procedure",
        )

        self.assertEqual(calls[0], (8, "hybrid", "none"))
        self.assertEqual(result["answer"], "Use EthicsPoint.")
        self.assertEqual(result["settings_used"]["retrieval_profile"], "auto")
        self.assertEqual(result["settings_used"]["resolved_retrieval_profile"], "auto_policy")
        self.assertTrue(result["settings_used"]["answer_verification"])

    def test_rag_pipeline_auto_multi_document_uses_broader_hybrid_search(self):
        calls = []
        pipeline = RAGPipeline.__new__(RAGPipeline)
        pipeline.retriever = type(
            "Retriever",
            (),
            {
                "retrieve": lambda self, query, top_k, method, reranker: calls.append(
                    (top_k, method, reranker)
                )
                or [
                    {
                        "filename": "support.pdf",
                        "page": 2,
                        "locator_label": "Page 2",
                        "chunk_id": 8,
                        "text": "Use the handbook for operational support guidance.",
                        "preview": "Use the handbook",
                        "rank": 1,
                        "score": 0.9,
                    }
                ]
            },
        )()
        pipeline.generator = type(
            "Generator",
            (),
            {"generate": lambda self, question, context, prompt_variant=None: "Use both sources."},
        )()

        result = pipeline.run(
            "How should we decide between product docs and the support handbook?",
            retrieval_profile="auto",
            question_type="multi_document_reasoning",
        )

        self.assertEqual(calls[0], (10, "hybrid", "none"))
        self.assertEqual(
            result["settings_used"]["resolved_retrieval_profile"],
            "auto_multi_document",
        )

    def test_rag_context_highlights_direct_requirement_sentences(self):
        context = _build_context(
            "What should company-wide announcements cover?",
            [
                {
                    "filename": "communication.pdf",
                    "page": 11,
                    "locator_label": "Page 11",
                    "text": (
                        "Common company wide announcements include organization changes. "
                        "Keep it simple, brief and summarize what is important. "
                        "Cover the 5 W's. What, Why, Who, When, Where."
                    ),
                }
            ],
        )

        high_signal = context.split("</high_signal_evidence>", 1)[0]
        self.assertIn("Cover the 5 W's", high_signal)
        self.assertIn("communication.pdf - Page 11", high_signal)
        self.assertLess(
            high_signal.index("Cover the 5 W's"),
            high_signal.index("Common company wide announcements include"),
        )

    def test_document_service_indexes_and_resets_documents(self):
        document = Document(
            filename="stored.pdf",
            original_filename="policy.pdf",
            file_path="stored.pdf",
            status="uploaded",
        )
        self.session.add(document)
        self.session.commit()

        service = DocumentService.__new__(DocumentService)
        service.db = self.session
        service.loader = type(
            "Loader",
            (),
            {
                "load_document": lambda self, path, original_filename=None: {
                    "document_type": "pdf",
                    "units": [
                        {
                            "unit_num": 1,
                            "page_num": 1,
                            "locator": "page:1",
                            "locator_label": "Page 1",
                            "section_title": None,
                            "sheet_name": None,
                            "text": "Refunds in 30 days",
                        }
                    ],
                    "total_pages": 1,
                    "total_units": 1,
                }
            },
        )()
        service.chunker = TextChunker()
        service.embeddings = type(
            "Embeddings", (), {"embed_texts": lambda self, texts: [[1.0, 0.0] for _ in texts]}
        )()
        service.vector_store = FakeVectorStore()

        result = service.index_documents(chunk_size=100, chunk_overlap=10)
        self.assertEqual(result["indexed_documents"], 1)
        self.assertEqual(document.status, "indexed")
        self.assertEqual(service.vector_store.upserts[0][1][0]["filename"], "policy.pdf")

        active = service.get_active_configuration()
        self.assertEqual(active, {"chunk_size": 100, "chunk_overlap": 10, "chunking_strategy": "auto"})
        self.assertTrue(service.reset_index())
        self.assertEqual(document.status, "uploaded")

    def test_index_configuration_change_marks_indexed_documents_stale(self):
        document = Document(
            filename="stored.pdf",
            original_filename="policy.pdf",
            file_path="stored.pdf",
            status="indexed",
            chunk_size=800,
            chunk_overlap=100,
            chunking_strategy="auto",
        )
        self.session.add(document)
        self.session.commit()

        service = DocumentService.__new__(DocumentService)
        service.db = self.session

        result = service.set_index_configuration(800, 150, "structure")

        self.assertTrue(result["changed"])
        self.assertTrue(result["reindex_required"])
        self.assertEqual(result["stale_documents"], 1)
        self.assertEqual(document.status, "needs_reindex")
        readiness = service.index_readiness()
        self.assertFalse(readiness["ready"])
        self.assertEqual(readiness["stale_documents"], 1)
        self.assertEqual(readiness["configuration"]["chunk_overlap"], 150)

    def test_document_service_retries_failed_documents_on_reindex(self):
        document = Document(
            filename="stored.pdf",
            original_filename="policy.pdf",
            file_path="stored.pdf",
            status="failed",
        )
        self.session.add(document)
        self.session.commit()

        service = DocumentService.__new__(DocumentService)
        service.db = self.session
        service.loader = type(
            "Loader",
            (),
            {
                "load_document": lambda self, path, original_filename=None: {
                    "document_type": "pdf",
                    "units": [
                        {
                            "unit_num": 1,
                            "page_num": 1,
                            "locator": "page:1",
                            "locator_label": "Page 1",
                            "section_title": None,
                            "sheet_name": None,
                            "text": "Recovered policy text",
                        }
                    ],
                    "total_pages": 1,
                    "total_units": 1,
                }
            },
        )()
        service.chunker = TextChunker()
        service.embeddings = type(
            "Embeddings", (), {"embed_texts": lambda self, texts: [[1.0, 0.0] for _ in texts]}
        )()
        service.vector_store = FakeVectorStore()

        result = service.index_documents(chunk_size=100, chunk_overlap=10)

        self.assertEqual(result["indexed_documents"], 1)
        self.assertEqual(document.status, "indexed")
        self.assertEqual(service.vector_store.upserts[0][1][0]["filename"], "policy.pdf")

    def test_document_upload_cleans_saved_file_when_pdf_loading_fails(self):
        service = DocumentService.__new__(DocumentService)
        service.db = self.session
        service.loader = type(
            "Loader", (), {"load_document": lambda self, path, original_filename=None: (_ for _ in ()).throw(ValueError("bad pdf"))}
        )()
        with patch(
            "app.services.document_service.save_upload_file",
            return_value=("stored.pdf", "stored.pdf"),
        ), patch("app.services.document_service.delete_upload_file") as delete_file:
            with self.assertRaisesRegex(Exception, "Failed to upload document"):
                service.upload_document("policy.pdf", b"%PDF-1.4")
        delete_file.assert_called_once_with("stored.pdf")


if __name__ == "__main__":
    unittest.main()
